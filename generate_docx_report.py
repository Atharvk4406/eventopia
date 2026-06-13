# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

OUT = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT, "SDG_EduAI_AtharvKudtarkar.docx")

doc = docx.Document()

# Set standard margins (1 inch on all sides)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper for cell background colors in tables
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Set base Normal style
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(55, 65, 81) # Slate Grey (#374151)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

# Formatting helpers
def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Navy (#1E3A8A)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(4, 120, 87) # Emerald (#047857)
    return p

def add_paragraph(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.bold = True
        r_pre.font.size = Pt(11)
        r_pre.font.color.rgb = RGBColor(31, 41, 55)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(55, 65, 81)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.bold = True
        r_pre.font.size = Pt(11)
        r_pre.font.color.rgb = RGBColor(31, 41, 55)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(55, 65, 81)
    return p

def create_table(headers, rows, col_widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A") # Navy
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            
    # Rows styling
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        for c_idx, cell_value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_value)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F3F4F6") # Light Grey alternating
            else:
                set_cell_background(cell, "FFFFFF")
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(31, 41, 55)
                run.font.name = 'Calibri'
                
    # Set widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)
            
    doc.add_paragraph() # Spacer after table

def add_image_block(filename, caption):
    path = os.path.join(OUT, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        p_cap.paragraph_format.keep_with_next = True
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Calibri'
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(100, 116, 139)
    else:
        print(f"Warning: Image {filename} not found in {OUT}")

# ──────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ──────────────────────────────────────────────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("PROJECT REPORT ON EDUAI")
run_title.font.name = 'Calibri'
run_title.font.size = Pt(26)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(30, 58, 138)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_subtitle.add_run("An AI-Powered Personalized Learning Companion\nFocus Area: United Nations SDG 4 – Quality Education")
run_sub.font.name = 'Calibri'
run_sub.font.size = Pt(13)
run_sub.font.color.rgb = RGBColor(100, 116, 139)

p_type = doc.add_paragraph()
p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_type.paragraph_format.space_before = Pt(20)
run_type = p_type.add_run("ARTIFICIAL INTELLIGENCE INTERNSHIP FINAL PROJECT")
run_type.font.name = 'Calibri'
run_type.font.size = Pt(11)
run_type.font.bold = True
run_type.font.color.rgb = RGBColor(4, 120, 87)

for _ in range(3):
    doc.add_paragraph()

headers = ["Project Metadata", "Details"]
rows = [
    ["Student Name", "Atharv Kudtarkar"],
    ["Project Title", "EduAI – Smart Learning Companion"],
    ["Sustainable Development Goal", "SDG 4 – Quality Education"],
    ["Artificial Intelligence Model", "Google Gemini 2.0 Flash (REST API)"],
    ["Application Platform", "HTML5 / CSS3 / Vanilla JavaScript"],
    ["Live Simulation URL", "http://eduai-sdg4.surge.sh"],
    ["Open Source Repository", "https://github.com/Atharvk4406/eventopia"],
    ["Submission Date", "June 2026"]
]
create_table(headers, rows, [2.5, 4.0])

# Add footer
footer = doc.sections[0].footer
f_p = footer.paragraphs[0]
f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
f_run = f_p.add_run("EduAI Final Project Report | Atharv Kudtarkar")
f_run.font.size = Pt(8.5)
f_run.font.color.rgb = RGBColor(148, 163, 184)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 1: SELECTED SDG AND REASON FOR SELECTION
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("1. Selected SDG and Reason for Selection")
add_heading_2("Selected Goal: United Nations SDG 4 – Quality Education")

add_paragraph(
    "The Sustainable Development Goal 4 (SDG 4) established by the United Nations aims to \"ensure inclusive "
    "and equitable quality education and promote lifelong learning opportunities for all\" by the year 2030. "
    "Education is recognized globally as a foundational driver of individual empowerment, socio-economic mobility, "
    "and national development. Inclusive and quality education can break cycles of poverty, reduce systemic "
    "inequality, and equip youth with the tools necessary to thrive in a rapidly changing knowledge economy."
)

add_paragraph(
    "However, modern educational systems face severe structural bottlenecks, particularly in developing regions. "
    "The choice of SDG 4 for this project is driven by several critical challenges that persist in traditional "
    "classroom environments today:"
)

add_bullet(
    "Overburdened Teachers: Student-to-teacher ratios in municipal and low-cost schools frequently exceed "
    "50:1. Under these conditions, teachers cannot provide personalized guidance or adjust their pacing to accommodate "
    "struggling students.", "High Student-to-Teacher Ratios: "
)
add_bullet(
    "Quality private tutoring and exam preparation resources are expensive, creating an unfair "
    "advantage for students from affluent families while leaving underprivileged students behind.", "Socio-Economic Barriers to Tutoring: "
)
add_bullet(
    "Traditional textbooks and pre-recorded online lectures are static. If a student is "
    "confused by a specific paragraph or step in a mathematical derivation, they cannot interact with the material to ask follow-up questions.", "Inefficiencies in Static Self-Study: "
)
add_bullet(
    "A self-study student studying late at night has no immediate recourse when they hit a concept they "
    "do not understand, leading to frustration, misconceptions, and decreased learning motivation.", "Lack of Real-Time Query Resolution: "
)

add_paragraph(
    "By leveraging generative AI, we can build tools that provide personalized, high-quality, and interactive tutoring "
    "at zero marginal cost. EduAI was developed to demonstrate how conversational AI can democratize quality "
    "learning, giving every student access to a patient, knowledgeable, and always-available private tutor."
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 2: PROBLEM STATEMENT
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("2. Problem Statement")

add_paragraph(
    "Despite high internet penetration and the availability of generic online courses, students still face major "
    "barriers when studying independently. The core problems that this project addresses are defined below:",
    "Problem Context: "
)

add_bullet(
    "A student studying a complex subject (e.g., quadratic equations in mathematics, "
    "photosynthesis in biology, or recursion in computer science) often experiences specific points of confusion. "
    "Without a personal tutor, there is no immediate way to get custom explanations tailored to their current level of "
    "understanding. This leads to gaps in knowledge that accumulate as the syllabus progresses.", "What is the problem? "
)
add_bullet(
    "This problem is prevalent in developing nations and semi-urban or rural school districts "
    "(such as public schools in India), where academic support infrastructure outside of school is practically non-existent "
    "and classrooms are heavily overcrowded.", "Where does the problem exist? "
)
add_bullet(
    "The primary group affected consists of high school and university students, particularly those from "
    "middle-income and lower-income families who must master technical subjects for board or competitive exams but cannot "
    "afford high-cost private coaching classes.", "Who is affected? "
)
add_bullet(
    "When academic difficulties are left unaddressed, students experience a gradual loss of academic self-esteem, "
    "which leads to poor performance in standard assessments and ultimately higher dropout rates in STEM streams. It prevents "
    "capable students from pursuing higher education simply because they lacked conceptual guidance at crucial learning stages.", "Why is the problem serious? "
)
add_bullet(
    "If educational support remains a commodity accessible only to the wealthy, the educational "
    "achievement gap between social classes will widen. Underprivileged students will be systemically locked out of high-paying "
    "technical and scientific jobs, reinforcing generational cycles of inequality.", "What could happen if it is not solved? "
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 3: PROPOSED SOLUTION
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("3. Proposed Solution")

add_paragraph(
    "EduAI is an interactive, web-based, AI-powered study companion designed to give students 24/7 access to "
    "structured, personalized tutoring. The system leverages the Google Gemini 2.0 Flash API to simulate an "
    "on-demand private educator, combining Q&A, adaptive assessment, study organization, and curated learning recommendations "
    "into a single free-to-use interface."
)

add_heading_2("How the System Works")
add_paragraph(
    "The user interface is built using client-side technologies (HTML5, CSS3, and Vanilla JavaScript) to ensure the application "
    "remains extremely lightweight, fast-loading, and responsive on mobile devices and low-bandwidth connections. The workflow operates as follows:"
)

add_bullet(
    "The student visits the live EduAI portal on any smartphone, tablet, or desktop computer. The "
    "entire experience is hosted statically and runs in the browser, requiring no registration or user logins.", "Accessing the Portal: "
)
add_bullet(
    "The student selects a target subject (e.g., Mathematics, Physics, Chemistry, Computer "
    "Science, English, or History) and enters their specific query or doubt in natural language.", "Interactive AI Tutoring: "
)
add_bullet(
    "The query is packaged with structured system instructions that tell the AI to act as a "
    "helpful, encouraging school teacher. The request is sent to the Google Gemini 2.0 Flash API, which returns a clear, step-by-step, "
    "and academically sound explanation. Mathematical formulas and code syntax are formatted for readability.", "Processing and AI Response: "
)
add_bullet(
    "To gauge their comprehension, the student can request an AI-generated multiple-choice quiz on "
    "any topic. The system parses the AI's JSON output to display an interactive quiz. When submitted, the application grades the answers "
    "instantly, highlighting correct selections in green and incorrect selections in red, alongside brief conceptual explanations.", "Adaptive Evaluation: "
)
add_bullet(
    "The student can log their tasks and track their progress using the client-side Goal Planner. "
    "As tasks are marked complete, a visual progress bar updates in real time. If the student completes at least one task a day, "
    "the local storage increments their study streak, providing gamified motivation to study daily.", "Structured Goal Tracking: "
)

add_heading_2("Mitigating the Problem and Core Audience")
add_paragraph(
    "EduAI addresses the core problems of self-study by removing the barrier of 'getting stuck'. It provides safe, non-judgmental, "
    "and immediate feedback. The primary audience for this application is secondary and higher-secondary school students "
    "(grades 9 through 12), undergraduate college students, and independent self-learners seeking a structured path to master new academic concepts."
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 4: PROJECT FEATURES
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("4. Project Features")

add_paragraph(
    "The EduAI platform is divided into five core functional areas designed to support the student through the entire learning cycle:"
)

add_bullet(
    "A conversational chat interface where students can ask complex questions and receive structured, "
    "step-by-step guidance. The AI tutor avoids giving direct answers to homework and instead guides the student through the logical steps "
    "to solve the problem themselves.", "AI Subject Tutor: "
)
add_bullet(
    "Generates dynamic 5-question multiple-choice quizzes on any user-specified topic. It evaluates "
    "answers dynamically, displays a final score card, and offers immediate explanations for each answer choice to ensure active learning.", "AI Quiz Generator: "
)
add_bullet(
    "A productivity checklist where students add, organize, and check off daily learning tasks. Features a "
    "visual progress bar showing percent completion and a study streak tracker that saves progress in localStorage.", "Study Planner & Streak Tracker: "
)
add_bullet(
    "Accessible through the tutor panel, this feature allows students to ask the AI for learning resource "
    "recommendations. The AI analyzes the subject and lists highly-rated textbooks, free interactive platforms, and educational YouTube channels.", "Resource Recommender: "
)
add_bullet(
    "The main landing page acts as a central hub, displaying live counters for total tutoring sessions, "
    "quizzes taken, goals completed, and current study streaks to give students a clear overview of their effort.", "Unified Student Dashboard: "
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 5: TECHNOLOGY STACK
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("5. Technology Stack")

add_paragraph(
    "The technology stack was chosen to prioritize high performance, portability, ease of deployment, and minimal operational costs. "
    "This allows the app to be run by a single developer at virtually zero hosting expense, making it highly scalable."
)

headers_tech = ["Stack Component", "Technology Used", "Role in Project"]
rows_tech = [
    ["Cognitive AI Engine", "Google Gemini 2.0 Flash (via REST API)", "Generates structured educational explanations, builds custom MCQ quizzes, and provides curated recommendations."],
    ["Frontend Structure", "HTML5 (Semantic Markup)", "Defines clean, accessible structures for the landing dashboard, tutoring chat, quiz panels, and planner layouts."],
    ["Design System", "Vanilla CSS3 (Glassmorphism & Flexbox/Grid)", "Applies a modern, dark-mode design system using CSS custom variables, backdrop filters, responsive grid structures, and transitions."],
    ["Client-Side Scripting", "Vanilla JavaScript (ES6+)", "Handles API fetch requests, parses JSON responses, dynamically updates DOM nodes, controls quiz grading, and manages local storage variables."],
    ["Client Data Persistence", "Web Storage API (localStorage)", "Saves student goals, completed task lists, tutoring logs, and study streaks locally on the user's browser, ensuring user privacy."],
    ["Hosting & Deployment", "Surge.sh (Global Static CDN)", "Hosts static frontend assets globally, providing fast load times and automatic HTTPS encryption."],
    ["Version Control", "Git & GitHub", "Tracks codebase changes and manages the open-source repository."]
]
create_table(headers_tech, rows_tech, [1.8, 2.2, 3.0])

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 6: PROJECT SCREENSHOTS (PROOF OF WORK)
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("6. Project Screenshots (Proof of Work)")
add_paragraph(
    "The screenshots below demonstrate the fully functional EduAI system running in a browser. Each interface "
    "is designed to be clean, intuitive, and responsive."
)

add_image_block(
    "p1_dashboard.png",
    "Figure 1: EduAI Main Dashboard. Displays active stats (sessions, quizzes, streaks), the SDG 4 mission summary, "
    "and quick navigation cards to launch the Tutor, Quiz, and Planner modules."
)

add_image_block(
    "p2_tutor.png",
    "Figure 2: AI Subject Tutor. Shows real-time conversations with the AI tutor resolving doubts in Mathematics "
    "(Pythagorean theorem), Computer Science (recursion), and History (WW2 causes)."
)

add_image_block(
    "p3_quiz.png",
    "Figure 3: AI Quiz Generator. Displays an interactive quiz on photosynthesis. Highlights selected, correct, "
    "and incorrect options with real-time grading and feedback."
)

add_image_block(
    "p4_planner.png",
    "Figure 4: Study Planner. Displays a daily study checklist with categorized subject tags, a visual percentage progress bar, "
    "and a streak tracker."
)

add_image_block(
    "p5_resources.png",
    "Figure 5: Resource Recommender. Demonstrates the AI recommendations page suggesting books, websites, "
    "and courses to learn Computer Science."
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 7: OPEN SOURCE REPOSITORY & DEPLOYMENT
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("7. Open Source Repository & Project Links")

add_paragraph(
    "The code and assets for this project are fully open source. The repository and live deployment links are as follows:"
)

add_bullet("https://github.com/Atharvk4406/eventopia", "GitHub Repository: ")
add_bullet("https://github.com/Atharvk4406/eventopia/tree/main/eduai", "EduAI Project Directory: ")
add_bullet("http://eduai-sdg4.surge.sh", "Live Interactive Application: ")

add_paragraph(
    "The open-source repository contains the following core files and directories that make up the client-side application:"
)

add_bullet("The primary landing page and user dashboard containing aggregated learning statistics.", "index.html: ")
add_bullet("The conversational chatbot tutor interface with specific subject selection dropdowns.", "tutor.html: ")
add_bullet("The interactive evaluation module containing the AI-generated quiz interface.", "quiz.html: ")
add_bullet("The daily checklist planner and streak tracking screen.", "planner.html: ")
add_bullet("Contains the glassmorphic dark theme styles, responsive flexbox layout definitions, and keyframe animations.", "style.css: ")
add_bullet("Contains the core logic for the Gemini API call handling, JSON query formatting, and localStorage CRUD operations.", "app.js: ")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 8: FUTURE SCOPE
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("8. Future Scope & Enhancements")

add_paragraph(
    "While EduAI is fully functional as a lightweight client-side application, several key improvements can be "
    "implemented in future releases to scale its impact:"
)

add_bullet(
    "Developing native mobile applications for Android and iOS using React Native or Flutter. This "
    "will allow push notifications for study reminders and support offline local database cache storage.", "Mobile Application: "
)
add_bullet(
    "Adding support for Indian regional languages (such as Hindi, Marathi, Telugu, and Tamil) to "
    "allow students in rural areas who are not proficient in English to leverage the AI tutor.", "Multi-Language Support: "
)
add_bullet(
    "Analyzing historical quiz scores and chat logs to automatically map a student's weak areas and "
    "recommend targeted practice sets, creating a truly adaptive study curriculum.", "Adaptive Learning Paths: "
)
add_bullet(
    "Creating a secondary login portal for teachers and parents to review students' study progress, "
    "assign specific topics to read, and monitor AI tutor chat histories for safety and guidance.", "Teacher & Parent Dashboard: "
)
add_bullet(
    "Integrating speech-to-text and text-to-speech APIs to allow students with visual impairments or "
    "reading difficulties to communicate with the tutor naturally using voice commands.", "Voice-Based Learning Interface: "
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 9: CONCLUSION & SUMMARY
# ──────────────────────────────────────────────────────────────────────────────
add_heading_1("9. Conclusion & Summary")

add_paragraph(
    "EduAI successfully demonstrates how Generative AI models like Google Gemini 2.0 Flash can be applied practically to "
    "support United Nations SDG 4 (Quality Education). By providing 24/7 personal tutoring, adaptive evaluations, structured "
    "organization, and resource recommendations, the application bridges the gap between passive content consumption and active "
    "conceptual learning."
)

add_paragraph(
    "Because the system runs entirely client-side and requires no dedicated backend servers, it operates at virtually zero cost, "
    "making it highly scalable and easy to maintain. Key learnings from building this project include masterfully integrating "
    "generative model APIs in raw JavaScript, designing a cohesive glassmorphic UI using modern CSS, managing browser-level "
    "data persistence, and implementing clean error handling to handle API request rate limits gracefully. The project stands as "
    "proof that accessible, impactful educational technology can be developed and deployed rapidly using modern AI APIs."
)

doc.save(DOCX_PATH)
print(f"Document successfully saved to {DOCX_PATH}")
